#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera o documento 'Catadores Digitais — Guia da Plataforma' em .docx."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paleta ──────────────────────────────────────────────────────────────
AZUL = RGBColor(0x00, 0x30, 0x87)       # azul institucional (mesmo do diploma/footer)
AZUL_CLARO = RGBColor(0x1E, 0x5C, 0xB3)
CINZA_TEXTO = RGBColor(0x33, 0x33, 0x33)
CINZA_CALLOUT = RGBColor(0x5A, 0x5A, 0x5A)
CINZA_CLARO_BG = "F2F2F7"
DOURADO = RGBColor(0xB8, 0x86, 0x0B)

doc = Document()

# ─── Estilos base ────────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = CINZA_TEXTO
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.25

for i, size, color, before, after in [(1, 22, AZUL, 28, 10), (2, 15, AZUL, 20, 6), (3, 12.5, AZUL_CLARO, 14, 4)]:
    st = doc.styles[f'Heading {i}']
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.page_break_before = False

sect = doc.sections[0]
sect.left_margin = Cm(2.5)
sect.right_margin = Cm(2.5)
sect.top_margin = Cm(2)
sect.bottom_margin = Cm(2)

# ─── Helpers ─────────────────────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def p(text, bold=False, italic=False, size=None, color=None, align=None, space_after=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        para.alignment = align
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    return para


def bullet(text, bold_lead=None):
    para = doc.add_paragraph(style='List Bullet')
    if bold_lead:
        r = para.add_run(bold_lead)
        r.bold = True
        para.add_run(text)
    else:
        para.add_run(text)
    return para


def numbered(items):
    for text in items:
        para = doc.add_paragraph(style='List Number')
        para.add_run(text)


def callout(text, label="Por trás dos panos"):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.6)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(12)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '9CA3AF')
    pBdr.append(left)
    pPr.append(pBdr)
    r1 = para.add_run(f"{label}: ")
    r1.bold = True
    r1.italic = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = CINZA_CALLOUT
    r2 = para.add_run(text)
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = CINZA_CALLOUT
    return para


def why_how_box(o_que, por_que, como):
    p("O que é", bold=True, size=11.5, color=AZUL_CLARO, space_after=2)
    p(o_que, space_after=10)
    p("Por que existe", bold=True, size=11.5, color=AZUL_CLARO, space_after=2)
    p(por_que, space_after=10)
    p("Como funciona", bold=True, size=11.5, color=AZUL_CLARO, space_after=2)
    if isinstance(como, list):
        numbered(como)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    else:
        p(como, space_after=10)


def simple_table(headers, rows, col_widths=None, header_color=AZUL):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr_cells[i].text = ''
        run = hdr_cells[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10.5)
        set_cell_shading(hdr_cells[i], '003087')
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(val)
            run.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def page_break():
    doc.add_page_break()


def hr():
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D1D5DB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_after = Pt(14)


# ═══════════════════════════════════════════════════════════════════════
# CAPA
# ═══════════════════════════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

p("CATADORES DIGITAIS", bold=True, size=34, color=AZUL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
p("Guia da Plataforma", size=20, color=CINZA_CALLOUT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
p("Como tudo funciona — e por que existe", italic=True, size=13, color=CINZA_CALLOUT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

p("Documento preparado para os times pedagógico e comercial", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
p("Programa de formação em tecnologia — Instituto Ipês, com apoio da Caixa (Fundo Socioambiental)", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
p("Agosto de 2026", size=10, color=CINZA_CALLOUT, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# SUMÁRIO
# ═══════════════════════════════════════════════════════════════════════

h1("Sumário")
secoes = [
    "1. Introdução",
    "2. Os três perfis de acesso",
    "3. A porta de entrada: a Landing Page",
    "4. Entrando na plataforma: login e segurança de acesso",
    "5. Turmas: o centro de tudo",
    "6. Aulas e o calendário",
    "7. Banco de Aulas: conteúdo pronto para reaproveitar",
    "8. Materiais de apoio",
    "9. Avaliações",
    "10. Frequência (chamada)",
    "11. Relatórios",
    "12. Diplomas e certificados",
    "13. Conquistas do aluno",
    "14. Anotações do aluno",
    "15. Resumo: o que cada perfil enxerga",
    "16. Segurança e privacidade dos dados",
    "17. O que vem a seguir",
]
for s in secoes:
    bullet(s)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 1. INTRODUÇÃO
# ═══════════════════════════════════════════════════════════════════════

h1("1. Introdução")

p(
    "A Catadores Digitais é um programa de formação gratuita em tecnologia, realizado pelo Instituto "
    "Ipês com apoio da Caixa, por meio do Fundo Socioambiental. O programa dá prioridade a catadoras e "
    "catadores de materiais recicláveis, seus familiares e moradores da Cidade Estrutural e entorno — "
    "pessoas que, historicamente, têm menos acesso a oportunidades de qualificação profissional em "
    "tecnologia."
)
p(
    "Este documento explica a plataforma digital que dá suporte a esse programa: o que cada parte dela "
    "faz, por que ela existe daquele jeito, e como as pessoas — alunos, professores e equipe "
    "administrativa — usam cada uma no dia a dia. A ideia é que qualquer pessoa dos times pedagógico e "
    "comercial consiga entender o sistema por completo, mesmo sem experiência técnica."
)

h2("Duas partes, um só programa")
p(
    "Existem, na prática, dois sistemas que trabalham juntos, com propósitos diferentes:"
)
bullet(
    " o site público do programa (catadoresdigitais.com.br). É a vitrine: qualquer pessoa pode "
    "visitar, entender o que é o programa, ver quem tem prioridade nas vagas e deixar seu contato para "
    "ser avisada quando as inscrições abrirem.",
    bold_lead="A Landing Page —",
)
bullet(
    " o sistema onde o curso realmente acontece — turmas, aulas, presença, materiais, avaliações e "
    "diplomas. Só é acessível a quem já faz parte do programa (aluno, professor ou administrador), "
    "com login.",
    bold_lead="A Plataforma —",
)
callout(
    "os dois foram construídos como projetos separados de propósito. A landing precisa ser leve e "
    "rápida para qualquer visitante — sem exigir login, sem carregar dados sensíveis. A plataforma "
    "interna lida com informações de alunos, presença e certificados, então tem uma estrutura mais "
    "robusta por trás, com controle de acesso em cada tela."
)

# ═══════════════════════════════════════════════════════════════════════
# 2. OS TRÊS PERFIS DE ACESSO
# ═══════════════════════════════════════════════════════════════════════

h1("2. Os três perfis de acesso")

p(
    "Toda pessoa que acessa a plataforma tem um destes três perfis, e o que ela enxerga e pode fazer "
    "muda de acordo com isso:"
)

simple_table(
    ["Perfil", "O que pode fazer"],
    [
        ("Administrador", "Acesso total: gerencia turmas, usuários, libera novos acessos, aprova aulas "
                           "criadas por professores, configura diplomas e vê relatórios de todas as turmas."),
        ("Professor", "Cria e edita aulas das turmas em que leciona, lança presença, cria avaliações, "
                       "gerencia materiais e pode emitir diplomas. Aulas que ele cria ficam pendentes de "
                       "aprovação do administrador antes de aparecerem oficialmente."),
        ("Aluno", "Acompanha suas turmas, acessa materiais e avaliações, confirma presença, acompanha sua "
                   "frequência, escreve anotações pessoais e vê seus diplomas conquistados."),
    ],
    col_widths=[Cm(3.2), Cm(11.3)],
)

h2("Quem entra na plataforma? A lista de liberação")
p(
    "O acesso não é livre — ninguém consegue simplesmente criar uma conta sozinho. Antes de qualquer "
    "pessoa conseguir entrar, um administrador precisa liberar o e-mail dela numa lista de acesso, "
    "informando se ela vai entrar como aluna/aluno ou professor(a), e em qual turma."
)
p(
    "Só depois desse cadastro prévio a pessoa consegue fazer login (com sua conta Google) e é "
    "automaticamente matriculada na turma certa, sem nenhum passo manual extra."
)
callout(
    "essa lista de liberação existe numa coleção de dados chamada allowlist. Sem estar nela (ou ser o "
    "e-mail do administrador principal), o login é recusado — é essa a camada que mantém a plataforma "
    "reservada só a quem realmente faz parte do programa."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 3. LANDING PAGE
# ═══════════════════════════════════════════════════════════════════════

h1("3. A porta de entrada: a Landing Page")

why_how_box(
    o_que=(
        "O site público do programa, hospedado em catadoresdigitais.com.br. Apresenta o programa, os "
        "cursos oferecidos, quem tem prioridade nas vagas (catadores, familiares e moradores da Cidade "
        "Estrutural) e um formulário para quem quer ser avisado quando as inscrições abrirem."
    ),
    por_que=(
        "Antes mesmo de abrir inscrições formalmente, o programa precisa de um jeito de se apresentar "
        "publicamente e começar a construir uma lista de pessoas realmente interessadas. Isso permite "
        "que os times pedagógico e comercial já comecem a se planejar — quantas pessoas demonstraram "
        "interesse, de onde vêm — antes mesmo da primeira turma abrir."
    ),
    como=[
        "A pessoa visitante navega pelo site e entende do que se trata o programa.",
        "Se tiver interesse, preenche o formulário “Seja o primeiro a saber” com o e-mail.",
        "Esse e-mail é salvo numa lista de interessados.",
        "O administrador acompanha essa lista dentro da plataforma, em Interessados da Landing, e pode "
        "usá-la para campanhas de divulgação, follow-up ou planejamento de novas turmas.",
    ],
)

h2("Proteção contra spam")
p(
    "Formulários públicos na internet costumam atrair tentativas automatizadas de cadastro em massa "
    "(spam). Para evitar isso sem exigir nenhum tipo de captcha incômodo para quem está preenchendo de "
    "boa fé, o formulário tem um mecanismo invisível que identifica e barra essas tentativas "
    "automatizadas."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 4. LOGIN E SEGURANÇA DE ACESSO
# ═══════════════════════════════════════════════════════════════════════

h1("4. Entrando na plataforma: login e segurança de acesso")

why_how_box(
    o_que=(
        "O login da plataforma é feito exclusivamente com conta Google — não existe usuário e senha "
        "próprios do sistema para criar, lembrar ou perder."
    ),
    por_que=(
        "Usar o login do Google significa que a plataforma nunca precisa guardar nem se preocupar com "
        "senhas — que é uma das formas mais comuns de vazamento de dados em qualquer sistema. Também "
        "torna o acesso mais simples: quem já usa Gmail no dia a dia entra com poucos cliques."
    ),
    como=(
        "A pessoa clica em “Entrar com Google”, confirma sua conta, e a plataforma verifica se "
        "aquele e-mail está na lista de liberação (allowlist). Se estiver, a conta é criada "
        "automaticamente com o papel (aluno ou professor) e a turma definidos por quem cadastrou o "
        "convite, e a pessoa já entra direto no seu painel. Se o e-mail do administrador principal for "
        "usado, essa conta sempre entra com acesso total, sem precisar de convite."
    ),
)
callout(
    "depois do login, a plataforma mantém a pessoa conectada com segurança através de um cookie de "
    "sessão — um jeito de “lembrar” que aquela pessoa já se autenticou, sem precisar repetir o "
    "processo a cada página visitada."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 5. TURMAS
# ═══════════════════════════════════════════════════════════════════════

h1("5. Turmas: o centro de tudo")

why_how_box(
    o_que=(
        "Uma turma representa uma edição de um curso — um grupo de alunos, professores, e um período "
        "definido (data de início e fim). É a unidade em torno da qual quase tudo na plataforma se "
        "organiza: aulas, materiais, avaliações, presença e diplomas pertencem sempre a uma turma "
        "específica."
    ),
    por_que=(
        "O programa pode ter várias turmas rodando ao mesmo tempo, em diferentes estágios, com "
        "diferentes públicos. Organizar tudo por turma evita que o conteúdo de uma turma se misture com "
        "o de outra, e permite acompanhar o progresso de cada grupo de forma independente."
    ),
    como=(
        "O administrador cria uma turma escolhendo nome, um ícone e uma cor (para identificação visual "
        "rápida em toda a plataforma), as datas de início e fim, e os alunos matriculados. Depois disso, "
        "a turma pode ser editada a qualquer momento, e cada aula, material e avaliação criados dentro "
        "dela ficam automaticamente vinculados àquela turma."
    ),
)

h2("O que acontece quando a turma termina")
p(
    "Quando a data final de uma turma passa, ela é automaticamente marcada como encerrada "
    "(“arquivada”). Isso não apaga nada — todo o histórico continua acessível para consulta e "
    "relatórios — mas passa a impedir edições acidentais: só o administrador consegue alterar uma turma "
    "já encerrada. Isso preserva o registro histórico do que de fato aconteceu naquele curso."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 6. AULAS E CALENDÁRIO
# ═══════════════════════════════════════════════════════════════════════

h1("6. Aulas e o calendário")

why_how_box(
    o_que=(
        "Cada turma tem um calendário mensal onde ficam marcadas as aulas: data, horário, título, "
        "descrição e os professores responsáveis. Cada aula tem sua própria página, com materiais, "
        "controle de presença e avaliações daquele encontro específico."
    ),
    por_que=(
        "Ter um calendário claro, visual e compartilhado evita o vaivém de mensagens perguntando "
        "“quando é a próxima aula?” e dá a alunos e professores uma visão imediata do que vem "
        "a seguir, direto na tela inicial de cada um."
    ),
    como=[
        "O professor ou administrador escolhe uma data no calendário e preenche título, horário e "
        "descrição da aula.",
        "Não é possível criar uma aula em uma data que já passou — isso evita erros de lançamento e "
        "mantém o calendário confiável.",
        "Quando é um professor quem cria a aula, ela entra como “pendente” até o administrador "
        "revisar e aprovar — uma camada extra de controle de qualidade antes de a aula ficar visível "
        "oficialmente para os alunos.",
        "A partir daí, a aula aparece no calendário e ganha sua própria página, onde entram materiais, "
        "presença e avaliações.",
    ],
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 7. BANCO DE AULAS
# ═══════════════════════════════════════════════════════════════════════

h1("7. Banco de Aulas: conteúdo pronto para reaproveitar")

why_how_box(
    o_que=(
        "Um espaço onde professores guardam aulas prontas — com título, descrição, materiais e "
        "avaliações — sem vincular a uma data específica. Quando for necessário, essa aula pronta pode "
        "ser “agendada” para uma data real em qualquer turma."
    ),
    por_que=(
        "Boa parte do conteúdo de um curso se repete de turma para turma — a introdução ao tema, os "
        "exercícios básicos, os materiais de referência. Sem o Banco de Aulas, o professor precisaria "
        "recriar essa aula do zero toda vez que uma nova turma chegasse ao mesmo ponto do curso. Com "
        "ele, o trabalho é feito uma vez e reaproveitado quantas vezes for preciso."
    ),
    como=(
        "O professor monta a aula no Banco (título, descrição, materiais, avaliações). Quando chega o "
        "momento de dar aquela aula numa turma específica, ele escolhe a data e horário e "
        "“agenda” — a aula é automaticamente colocada no calendário daquela turma, já pronta. "
        "A mesma aula do banco pode ser agendada de novo, em outra data ou outra turma, sem duplicar "
        "trabalho."
    ),
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 8. MATERIAIS DE APOIO
# ═══════════════════════════════════════════════════════════════════════

h1("8. Materiais de apoio")

why_how_box(
    o_que=(
        "Dentro de cada aula, professores podem adicionar links (vídeos do YouTube, arquivos do Google "
        "Drive, apresentações, documentos) e blocos de texto livre — tudo o que o aluno precisa para "
        "acompanhar aquele conteúdo."
    ),
    por_que=(
        "Centraliza tudo o que o aluno precisa num único lugar, ligado diretamente à aula certa — sem "
        "depender de grupos de WhatsApp, e-mails avulsos ou pastas compartilhadas soltas, onde material "
        "se perde com facilidade."
    ),
    como=(
        "A plataforma reconhece automaticamente se um link é um vídeo ou um documento, e ajusta como "
        "ele é exibido. Blocos de texto podem ser expandidos direto na tela, sem precisar abrir uma "
        "janela nova. Materiais podem ser reordenados a qualquer momento, para manter a sequência que "
        "faz mais sentido pedagogicamente."
    ),
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 9. AVALIAÇÕES
# ═══════════════════════════════════════════════════════════════════════

h1("9. Avaliações")

why_how_box(
    o_que=(
        "Cada aula pode ter uma ou mais avaliações, em três formatos: link (o aluno envia o link de um "
        "trabalho realizado em outra ferramenta), texto (resposta aberta, escrita) ou quiz (múltipla "
        "escolha, com correção automática)."
    ),
    por_que=(
        "Formatos diferentes servem a propósitos pedagógicos diferentes — um quiz é rápido de aplicar e "
        "corrigir, uma resposta em texto permite avaliar raciocínio e expressão, e o formato de link "
        "permite avaliar trabalhos feitos em ferramentas externas (como um site ou design criado pelo "
        "aluno)."
    ),
    como=[
        "O professor cria a avaliação escolhendo o formato e o conteúdo (pergunta, opções, etc).",
        "Antes de publicar, o professor pode usar “Testar avaliação” — uma simulação exata de "
        "como o aluno vai ver e responder, para conferir se está tudo certo.",
        "O aluno vê um botão simples de “Responder Avaliação”. Depois de enviar, um indicador "
        "mostra quantas perguntas já foram respondidas, e a avaliação não pode ser reaberta — o que "
        "preserva a integridade do resultado.",
        "Quizzes são corrigidos automaticamente, e o professor vê um indicador de certo/errado por "
        "aluno. Respostas em texto e link ficam disponíveis para leitura do professor.",
    ],
)
callout(
    "hoje a plataforma ainda não tem um recurso de correção/comentário do professor sobre respostas "
    "abertas (texto ou link) — ele consegue lê-las, mas não devolver um retorno escrito dentro do "
    "sistema. Está listado como próximo passo."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 10. FREQUÊNCIA
# ═══════════════════════════════════════════════════════════════════════

h1("10. Frequência (chamada)")

why_how_box(
    o_que=(
        "O controle de presença de cada aula. Existem dois jeitos de registrar quem esteve presente: um "
        "código de check-in que o próprio aluno usa durante a aula, e um lançamento manual feito pelo "
        "professor."
    ),
    por_que=(
        "Acompanhar presença é essencial num programa de formação — tanto para dar suporte pedagógico "
        "(identificar cedo quem está se afastando do curso) quanto para prestar contas aos parceiros "
        "que apoiam o programa (Instituto Ipês e Caixa)."
    ),
    como=[
        "Durante a aula, o professor revela um código de 4 dígitos.",
        "Cada aluno presente digita esse código na página da aula, pelo próprio celular ou computador — "
        "como uma chamada de sala de aula, só que instantânea e sem depender de o professor chamar nome "
        "por nome.",
        "Depois da aula, o professor também pode ajustar manualmente quem estava presente ou faltou, "
        "caso necessário.",
    ],
)

h2("Acompanhamento e alerta")
p(
    "Cada aluno vê sua própria porcentagem de frequência no painel principal. Se ela cair abaixo de "
    "85%, um aviso aparece incentivando o aluno a procurar o professor ou a coordenação — uma forma de "
    "agir antes que o afastamento se torne evasão."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 11. RELATÓRIOS
# ═══════════════════════════════════════════════════════════════════════

h1("11. Relatórios")

why_how_box(
    o_que=(
        "Um relatório por turma, disponível para administrador e professor, com uma tabela cruzando "
        "cada aluno com cada aula (presença e conclusão), um resumo da presença média da turma, filtro "
        "por período e exportação em CSV (planilha)."
    ),
    por_que=(
        "Dados de presença e conclusão são a base para dois tipos de decisão: pedagógica (onde a turma "
        "está tendo dificuldade, quem precisa de atenção) e institucional (prestação de contas para "
        "quem apoia o programa, com dados concretos e exportáveis)."
    ),
    como=(
        "Basta abrir o relatório da turma desejada. A tabela já vem pronta, e o filtro por período "
        "permite focar num recorte específico (por exemplo, só o último mês). O botão de download gera "
        "uma planilha para uso em apresentações, prestações de conta ou análises mais aprofundadas fora "
        "da plataforma."
    ),
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 12. DIPLOMAS
# ═══════════════════════════════════════════════════════════════════════

h1("12. Diplomas e certificados")

p(
    "Essa é uma das funcionalidades mais recentes da plataforma, e vale um detalhamento maior — ela foi "
    "pensada para reconhecer conquistas dos alunos de um jeito mais flexível do que o tradicional "
    "“certificado só no fim do curso”."
)

h2("O problema que ela resolve")
p(
    "Cursos mais longos costumam ter etapas relevantes ao longo do caminho — a conclusão de um módulo, "
    "por exemplo — que merecem reconhecimento próprio, mesmo antes do fim do curso inteiro. Esperar até "
    "a formatura final para reconhecer qualquer conquista significa perder uma ferramenta importante de "
    "motivação, e também deixa sem documento formal quem concluiu uma parte significativa do curso mas, "
    "por algum motivo, não chegou ao final."
)

h2("Como funciona, passo a passo")
numbered([
    "Cada turma tem um coordenador ou coordenadora geral configurado — nome e assinatura — que é a "
    "pessoa que “assina” os diplomas daquela turma.",
    "O administrador ou professor cria um marco: um ponto específico do curso que merece um diploma "
    "(por exemplo, “Conclusão do Módulo 1”), com título, data em que foi alcançado, carga "
    "horária correspondente, e a lista de alunos que chegaram até ali.",
    "Emitir os diplomas é uma etapa separada e deliberada, feita depois — dá tempo para revisar a lista "
    "de alunos antes de gerar os documentos definitivos.",
    "Ao emitir, a plataforma gera um diploma individual em PDF para cada aluno selecionado, pronto para "
    "download e impressão.",
])

h2("O que tem no diploma")
bullet("Logo do Instituto Ipês e logo da Caixa (Fundo Socioambiental) — os parceiros do programa.")
bullet("Nome completo do aluno e CPF.")
bullet("Título da conquista, carga horária e data.")
bullet("Nome da turma / curso.")
bullet("Assinatura do coordenador geral responsável.")
bullet("Um QR code de verificação.")

h2("O QR code de verificação")
p(
    "Cada diploma vem com um QR code que, ao ser escaneado, leva a uma página pública da plataforma "
    "confirmando que aquele diploma é autêntico — com os mesmos dados impressos no papel. Isso dá "
    "credibilidade extra ao documento: um empregador, por exemplo, pode verificar a autenticidade em "
    "segundos, sem precisar entrar em contato com o Instituto."
)

h2("Diplomas não mudam depois de emitidos")
p(
    "Uma vez emitido, um diploma é definitivo — funciona como uma fotografia do momento em que foi "
    "gerado. Mesmo que, depois, o coordenador da turma seja trocado ou algum dado seja atualizado, os "
    "diplomas já entregues permanecem exatamente como foram emitidos. É assim que um certificado "
    "deveria se comportar: uma vez emitido, não muda retroativamente."
)
callout(
    "isso significa também que, se um marco for excluído por engano ou revisão, os diplomas já emitidos "
    "a partir dele não são afetados nem apagados — só deixam de existir novos marcos daquele tipo."
)

h2("O CPF do aluno")
p(
    "O CPF é uma informação obrigatória no diploma, mas nem todo aluno tem esse dado cadastrado no "
    "perfil dele. Quando isso acontece, a própria tela de emissão pede o CPF na hora, e ele é gravado "
    "apenas naquele diploma específico — sem depender de o aluno ter preenchido o cadastro completo "
    "antes."
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 13. CONQUISTAS DO ALUNO
# ═══════════════════════════════════════════════════════════════════════

h1("13. Conquistas do aluno")

why_how_box(
    o_que=(
        "Uma aba, dentro da tela de cada turma, onde o aluno vê seus próprios diplomas — os que já "
        "conquistou e os marcos que ainda estão “em andamento” (ele já foi selecionado para "
        "aquele diploma, mas ainda não foi emitido oficialmente)."
    ),
    por_que=(
        "Dá visibilidade e motivação: o aluno enxerga o próprio progresso de forma concreta, e sabe "
        "exatamente o que falta para a próxima conquista, em vez de descobrir tudo de surpresa apenas no "
        "fim do curso."
    ),
    como=(
        "Cada diploma já conquistado aparece com um link direto para a página de verificação, de onde "
        "também é possível baixar o PDF. Os marcos ainda pendentes aparecem sinalizados como "
        "“aguardando emissão”."
    ),
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 14. ANOTAÇÕES DO ALUNO
# ═══════════════════════════════════════════════════════════════════════

h1("14. Anotações do aluno")

why_how_box(
    o_que=(
        "Um espaço pessoal, dentro da tela da turma, onde cada aluno pode escrever suas próprias "
        "anotações — texto livre, com formatação simples (negrito, listas, títulos)."
    ),
    por_que=(
        "Incentiva um hábito de estudo mais ativo: em vez de só consumir o conteúdo da aula, o aluno tem "
        "um lugar próprio, dentro da mesma plataforma, para registrar o que está aprendendo, sem "
        "precisar abrir outro aplicativo."
    ),
    como=(
        "As anotações são salvas automaticamente enquanto o aluno escreve, e são visíveis apenas para "
        "ele mesmo — nem outros alunos, nem professores, têm acesso a esse conteúdo pessoal."
    ),
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 15. RESUMO POR PERFIL
# ═══════════════════════════════════════════════════════════════════════

h1("15. Resumo: o que cada perfil enxerga")

p(
    "Dentro da tela de uma turma, a plataforma se organiza em abas. Nem toda aba aparece para todo "
    "mundo — a tabela abaixo resume quem vê o quê."
)

simple_table(
    ["Aba", "Administrador", "Professor", "Aluno"],
    [
        ("Visão geral (estatísticas da turma)", "Sim", "Não", "Não"),
        ("Conteúdo (aulas e materiais)", "Sim", "Sim", "Sim"),
        ("Presenças", "Sim", "Sim", "Não"),
        ("Professores", "Sim", "Sim", "Sim (dados limitados)"),
        ("Banco de Aulas", "Sim", "Sim", "Não"),
        ("Diplomas (criar/emitir)", "Sim", "Sim", "Não"),
        ("Conquistas (ver os próprios diplomas)", "Não", "Não", "Sim"),
        ("Anotações", "Não", "Não", "Sim"),
    ],
    col_widths=[Cm(5.6), Cm(3), Cm(2.7), Cm(3.2)],
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 16. SEGURANÇA E PRIVACIDADE
# ═══════════════════════════════════════════════════════════════════════

h1("16. Segurança e privacidade dos dados")

bullet(
    " ninguém entra na plataforma sem antes ter o e-mail liberado por um administrador — não existe "
    "cadastro aberto ao público em geral.",
    bold_lead="Acesso controlado —",
)
bullet(
    " como o login é feito pela conta Google de cada pessoa, a plataforma nunca guarda nem gerencia "
    "senhas próprias.",
    bold_lead="Sem senhas para vazar —",
)
bullet(
    " os dados de alunos, turmas, presença e diplomas ficam guardados em infraestrutura de nuvem do "
    "Google (a mesma tecnologia por trás de serviços como Gmail e Google Drive), com controle de acesso "
    "em cada tipo de informação.",
    bold_lead="Dados guardados com segurança —",
)
bullet(
    " a única informação pensada para ser pública é a página de verificação de um diploma específico — "
    "e mesmo essa só é alcançável por quem tem o link exato (como um link de compartilhamento do Google "
    "Docs). Nenhuma lista de alunos, presença ou dado pessoal fica exposta sem login.",
    bold_lead="Diplomas são a única informação pública —",
)
bullet(
    " o CPF de um aluno só é solicitado no momento em que é realmente necessário (emissão de um "
    "diploma), e cada pessoa só edita os próprios dados pessoais.",
    bold_lead="CPF tratado com cuidado —",
)

page_break()

# ═══════════════════════════════════════════════════════════════════════
# 17. O QUE VEM A SEGUIR
# ═══════════════════════════════════════════════════════════════════════

h1("17. O que vem a seguir")

p(
    "Nenhuma plataforma nasce pronta — a lista abaixo reúne os próximos passos já identificados, na "
    "ordem em que fazem mais sentido:"
)
numbered([
    "Permitir matricular e desmatricular alunos e professores direto na tela de edição da turma — hoje "
    "isso só é possível pelo perfil individual de cada usuário, no painel do administrador.",
    "Migrar a plataforma para o endereço definitivo (www.catadoresdigitais.com.br/plataforma) — hoje "
    "ela roda num endereço provisório enquanto essa configuração de domínio é finalizada.",
    "Completar a carga horária de alguns diplomas de teste criados antes de esse campo existir na "
    "plataforma.",
    "Permitir que o administrador edite o CPF de outra pessoa diretamente — hoje cada pessoa só edita o "
    "próprio CPF no seu perfil.",
    "Finalizar o SEO da Landing Page (como o site aparece quando compartilhado ou buscado no Google), "
    "para ampliar o alcance da divulgação.",
])

p("")
hr()
p(
    "Este documento reflete o estado da plataforma em agosto de 2026 e deve ser atualizado conforme "
    "novas funcionalidades forem lançadas.",
    italic=True, size=10, color=CINZA_CALLOUT,
)

doc.save('/Users/heliodev/Desktop/catadoresDigitais/docs/Catadores-Digitais-Guia-da-Plataforma.docx')
print("Documento gerado com sucesso.")
